"""Apply BID check-in feedback to cheat sheets for Sessions 1, 2, and 5."""

import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from copy import deepcopy
from lxml import etree
from docx.oxml.ns import qn


def insert_paragraph_before(ref_paragraph, text, bold=False, italic=False, size=None, color=None):
    """Insert a new paragraph before ref_paragraph in the document XML."""
    new_p = deepcopy(ref_paragraph._element)
    # Clear existing content runs
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)

    # Remove heading style so inserted text is body text
    pPr = new_p.find(qn("w:pPr"))
    if pPr is not None:
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is not None:
            pPr.remove(pStyle)
        # Remove numbering
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)

    # Add run with text
    r = etree.SubElement(new_p, qn("w:r"))
    rPr = etree.SubElement(r, qn("w:rPr"))
    if bold:
        etree.SubElement(rPr, qn("w:b"))
    if italic:
        etree.SubElement(rPr, qn("w:i"))
    if size:
        sz = etree.SubElement(rPr, qn("w:sz"))
        sz.set(qn("w:val"), str(size * 2))
    if color:
        c = etree.SubElement(rPr, qn("w:color"))
        c.set(qn("w:val"), color)

    t = etree.SubElement(r, qn("w:t"))
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    ref_paragraph._element.addprevious(new_p)
    return new_p


def insert_block(doc, ref_para, items):
    """Insert a list of (text, bold, italic, size, color) tuples before ref_para, in order."""
    for text, bold, italic, size, color in reversed(items):
        insert_paragraph_before(ref_para, text, bold=bold, italic=italic, size=size, color=color)


# ================================================================
# SESSION 1 — Pedagogical rationale + bridge
# ================================================================
doc1 = Document("case_study/session_materials/06_CaseStudy_CheatSheet_Session1.docx")

ex3_idx = None
for i, p in enumerate(doc1.paragraphs):
    if "EXERCISE 3: CONSTRUCT CAUSAL PATHWAY" in p.text:
        ex3_idx = i
        break

print(f"Session 1: Exercise 3 heading at paragraph {ex3_idx}")
ref1 = doc1.paragraphs[ex3_idx]

block1 = [
    ("", False, False, None, None),
    ("PEDAGOGICAL RATIONALE: WHY ORGANIZE BEHAVIORS BY ACTOR?", True, False, 11, "1F4E79"),
    ("", False, False, None, None),
    (
        "Organizing behaviors by actor ensures that the Theory of Change specifies "
        "whose behavior must change and what each actor must do differently. Without "
        "this, ToCs remain vague (\"capacity is built\") and untestable. Actor-specific "
        "behaviors: (1) make the ToC evaluable \u2014 each behavior can be measured with "
        "a specific indicator; (2) reveal the indirect pathway \u2014 the program reaches "
        "facilitators and caregivers, but the ultimate target is the child; (3) expose "
        "dependencies \u2014 facilitator quality affects caregiver capability, which affects "
        "child outcomes; (4) enable contribution analysis \u2014 if outcomes improve but "
        "specific behaviors did not change, the program's contribution claim weakens. "
        "In short: behaviors organized by actor turn the ToC from a diagram into a "
        "testable causal argument.",
        False, False, None, None,
    ),
    ("", False, False, None, None),
    ("BRIDGE TO EXERCISE 3: FROM BEHAVIORS TO CAUSAL CHAIN", True, False, 11, "1F4E79"),
    ("", False, False, None, None),
    (
        "In Exercise 3, participants will take these actor-specific behaviors and "
        "arrange them into a single causal pathway \u2014 connecting program outputs to "
        "behavioral changes to child outcomes. The behaviors defined above become the "
        "arrows in the causal chain: each one is a testable link. If a behavior does "
        "not occur, the chain breaks at that point. This is why specifying behaviors "
        "precisely in Exercise 2 is essential before constructing the pathway in "
        "Exercise 3.",
        False, True, None, None,
    ),
    ("", False, False, None, None),
]

insert_block(doc1, ref1, block1)
doc1.save("case_study/session_materials/06_CaseStudy_CheatSheet_Session1.docx")
print("Session 1 cheat sheet saved.")


# ================================================================
# SESSION 2 — Qualitative indicators section
# ================================================================
doc2 = Document("case_study/session_materials/07_CaseStudy_CheatSheet_Session2.docx")

takeaway_idx = None
for i, p in enumerate(doc2.paragraphs):
    if "KEY TAKEAWAYS" in p.text:
        takeaway_idx = i
        break

print(f"Session 2: KEY TAKEAWAYS at paragraph {takeaway_idx}")
ref2 = doc2.paragraphs[takeaway_idx]

block2 = [
    ("", False, False, None, None),
    ("QUALITATIVE INDICATORS IN THE RESULTS MATRIX", True, False, 11, "1F4E79"),
    ("", False, False, None, None),
    (
        "Not all indicators are numeric. Qualitative indicators capture dimensions "
        "of change that numbers alone cannot \u2014 depth of understanding, quality of "
        "interaction, fidelity of delivery, and shifts in beliefs or identity. "
        "Integrating qualitative indicators into the Results Matrix strengthens the "
        "contribution claim by providing explanatory evidence alongside measurement.",
        False, False, None, None,
    ),
    ("", False, False, None, None),
    ("TYPES OF QUALITATIVE INDICATORS:", True, False, None, None),
    ("", False, False, None, None),
    (
        "1. Observation rubrics: Structured rating scales applied by trained observers "
        "to assess quality of behaviors (e.g., quality of caregiver-child interaction "
        "rated on a 4-point rubric: no interaction / passive presence / responsive "
        "engagement / scaffolded learning).",
        False, False, None, None,
    ),
    ("", False, False, None, None),
    (
        "2. Fidelity scores: Composite scores assessing whether program delivery "
        "matches the intended design (e.g., session fidelity checklist: facilitator "
        "covers all 5 core topics, uses at least 2 participatory techniques, allows "
        "practice time \u2014 score 0\u20135).",
        False, False, None, None,
    ),
    ("", False, False, None, None),
    (
        "3. Narrative change indicators: Participant self-reports of change captured "
        "through structured prompts (e.g., Most Significant Change stories coded by "
        "theme; proportion of caregivers spontaneously mentioning new parenting identity).",
        False, False, None, None,
    ),
    ("", False, False, None, None),
    (
        "4. Thematic saturation indicators: Frequency with which key themes emerge "
        "in qualitative data collection (e.g., \"moments of realization\" reported by "
        "80% of interviewed caregivers across 6 municipalities).",
        False, False, None, None,
    ),
    ("", False, False, None, None),
    ("EXAMPLE: RESULTS MATRIX ROW WITH QUALITATIVE INDICATOR", True, False, 11, "1F4E79"),
    ("", False, False, None, None),
    ("Results Matrix Row \u2014 Observation Rubric:", True, False, None, None),
    ("NAME: Quality of caregiver-child interaction (observation rubric)", False, False, None, None),
    ("LEVEL: Behavior change \u2014 Observable behavior (qualitative)", False, False, None, None),
    (
        "DESCRIPTION: Trained observer rates caregiver-child interaction during 15-min "
        "structured play session using 4-point rubric (1=No interaction, 2=Passive "
        "presence, 3=Responsive engagement, 4=Scaffolded learning)",
        False, False, None, None,
    ),
    ("FORMULA: Mean rubric score across observed sessions; % scoring 3 or 4", False, False, None, None),
    ("SOURCE: Structured observation by trained home visitors (inter-rater reliability > 0.80)", False, False, None, None),
    ("BASELINE: Mean score 1.8; 20% scoring 3+", False, False, None, None),
    ("TARGET M12: Mean score 3.2; 65% scoring 3+ (treatment); Mean 2.0; 25% scoring 3+ (comparison)", False, False, None, None),
    ("FREQUENCY: Months 3, 6, 9, 12 (subsample of 200 dyads per group)", False, False, None, None),
    ("", False, False, None, None),
    ("Results Matrix Row \u2014 Fidelity Score:", True, False, None, None),
    ("NAME: Session delivery fidelity score", False, False, None, None),
    ("LEVEL: Output quality (qualitative)", False, False, None, None),
    (
        "DESCRIPTION: Composite score (0\u20135) based on checklist: facilitator covers core "
        "topics (1pt), uses participatory techniques (1pt), allows caregiver practice "
        "(1pt), provides individual feedback (1pt), manages time within 10% of plan (1pt)",
        False, False, None, None,
    ),
    ("FORMULA: Mean fidelity score across observed sessions; % scoring 4+", False, False, None, None),
    ("SOURCE: Supervisor observation using standardized checklist (monthly for 20% of sessions)", False, False, None, None),
    ("BASELINE: Not applicable (new program)", False, False, None, None),
    ("TARGET: Mean 4.0+ by Month 6; 80% of sessions scoring 4+", False, False, None, None),
    ("FREQUENCY: Monthly supervisor visits; quarterly external audit", False, False, None, None),
    ("", False, False, None, None),
    ("HOW QUALITATIVE AND QUANTITATIVE INDICATORS WORK TOGETHER:", True, False, None, "1F4E79"),
    ("", False, False, None, None),
    (
        "Quantitative indicator (interaction time): Tells us HOW MUCH interaction "
        "changed (15\u219245 min).",
        False, False, None, None,
    ),
    (
        "Qualitative indicator (observation rubric): Tells us HOW WELL the interaction "
        "quality changed (passive\u2192responsive).",
        False, False, None, None,
    ),
    (
        "Together: A caregiver who increases time but stays passive (high quantity, "
        "low quality) has NOT achieved the mechanism. Both indicators are needed to "
        "test the ToC.",
        False, True, None, None,
    ),
    ("", False, False, None, None),
]

insert_block(doc2, ref2, block2)
doc2.save("case_study/session_materials/07_CaseStudy_CheatSheet_Session2.docx")
print("Session 2 cheat sheet saved.")


# ================================================================
# SESSION 5 — Templates for Exercises 3 and 4
# ================================================================
doc5 = Document("case_study/session_materials/10_CaseStudy_CheatSheet_Session5.docx")

# Find Exercise 3 answer key (first "ANSWER KEY:" after Exercise 3 heading)
ex3_heading_idx = None
ex3_answer_idx = None
for i, p in enumerate(doc5.paragraphs):
    if "EXERCISE 3: ASSESS ASSUMPTION" in p.text:
        ex3_heading_idx = i
    if ex3_heading_idx is not None and "ANSWER KEY:" in p.text and ex3_answer_idx is None and i > ex3_heading_idx:
        ex3_answer_idx = i
        break

print(f"Session 5: Ex3 answer key at paragraph {ex3_answer_idx}")

SEP = "\u2500" * 73

assumptions = [
    "Facilitators can be recruited, trained, and retained",
    "Municipalities provide space and logistical support",
    "Health centers have capacity to add services",
    "Caregivers attend regularly (\u226570% of sessions)",
    "Knowledge/skills translate to daily practice",
    "Changed behavior sustains after program ends",
    "Dosage and quality of changed interaction is sufficient",
    "No overwhelming adversities swamp parenting improvements",
]

ex3_template = [
    ("", False, False, None, None),
    ("PARTICIPANT TEMPLATE \u2014 ASSUMPTION ASSESSMENT TABLE", True, False, 11, "1F4E79"),
    ("", False, False, None, None),
    (
        "Instructions: Complete one row per critical assumption from Session 1. "
        "For each assumption, summarize the quantitative and qualitative evidence, "
        "assess its status, and state the implication for the contribution claim.",
        False, True, None, None,
    ),
    ("", False, False, None, None),
]

for idx, assumption in enumerate(assumptions, 1):
    ex3_template.extend([
        (SEP, False, False, None, "999999"),
        (f"ASSUMPTION {idx}: {assumption}", True, False, None, None),
        ("Evidence (Quantitative): _______________________________________________", False, False, None, None),
        ("Evidence (Qualitative): ________________________________________________", False, False, None, None),
        ("Status:  [ ] Held   [ ] Partially held   [ ] Violated", False, False, None, None),
        ("Implication for contribution claim: ____________________________________", False, False, None, None),
    ])

ex3_template.extend([
    (SEP, False, False, None, "999999"),
    ("", False, False, None, None),
    ("SYNTHESIS: Overall, how many assumptions held / partially held / violated?", True, False, None, None),
    ("What pattern emerges? Where is the ToC strongest? Where is it weakest?", False, True, None, None),
    ("___________________________________________________________________________", False, False, None, None),
    ("___________________________________________________________________________", False, False, None, None),
    ("", False, False, None, None),
])

insert_block(doc5, doc5.paragraphs[ex3_answer_idx], ex3_template)

# Find Exercise 4 answer key
ex4_answer_idx = None
for i, p in enumerate(doc5.paragraphs):
    if "ANSWER KEY: STRONG CONTRIBUTION CLAIM" in p.text:
        ex4_answer_idx = i
        break

print(f"Session 5: Ex4 answer key at paragraph {ex4_answer_idx}")

criteria = [
    ("CONSISTENCY", "Do all indicators (outputs, mechanisms, behaviors, outcomes) point in the same direction?"),
    ("SPECIFICITY", "Are the changes specific to the expected mechanisms, or could any intervention have produced them?"),
    ("DOSE-RESPONSE", "Does more program exposure lead to better outcomes? Is there a gradient?"),
    ("TIMELINE", "Do changes follow the expected temporal sequence (capability first, then behavior, then outcomes)?"),
    ("ALTERNATIVES", "Have alternative explanations been systematically examined and ruled out (or accounted for)?"),
]

ex4_template = [
    ("", False, False, None, None),
    ("PARTICIPANT TEMPLATE \u2014 CONTRIBUTION CLAIM RATING TABLE", True, False, 11, "1F4E79"),
    ("", False, False, None, None),
    (
        "Instructions: For each criterion, answer the guiding question using evidence "
        "from Exercises 1\u20133. Summarize the evidence, then rate the criterion. After "
        "completing all five rows, assign an overall contribution claim rating.",
        False, True, None, None,
    ),
    ("", False, False, None, None),
]

for idx, (name, question) in enumerate(criteria, 1):
    ex4_template.extend([
        (SEP, False, False, None, "999999"),
        (f"CRITERION {idx}: {name}", True, False, None, None),
        (f"Guiding question: {question}", False, False, None, None),
        ("Evidence summary: ______________________________________________________", False, False, None, None),
        ("________________________________________________________________________", False, False, None, None),
        ("Rating:  [ ] Weak   [ ] Moderate   [ ] Strong   [ ] Very Strong", False, False, None, None),
    ])

ex4_template.extend([
    (SEP, False, False, None, "999999"),
    ("", False, False, None, None),
    ("OVERALL CONTRIBUTION CLAIM RATING", True, False, 12, "1F4E79"),
    ("", False, False, None, None),
    ("Rating:  [ ] Weak   [ ] Moderate   [ ] Strong   [ ] Very Strong", True, False, None, None),
    ("", False, False, None, None),
    ("Justification (2\u20133 sentences): _________________________________________", False, False, None, None),
    ("________________________________________________________________________", False, False, None, None),
    ("________________________________________________________________________", False, False, None, None),
    ("", False, False, None, None),
    ("Key strength of the claim: _____________________________________________", False, False, None, None),
    ("Main limitation or caveat: _____________________________________________", False, False, None, None),
    ("", False, False, None, None),
])

insert_block(doc5, doc5.paragraphs[ex4_answer_idx], ex4_template)

doc5.save("case_study/session_materials/10_CaseStudy_CheatSheet_Session5.docx")
print("Session 5 cheat sheet saved.")

print("\nAll three cheat sheets updated successfully.")
